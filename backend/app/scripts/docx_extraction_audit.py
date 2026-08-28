"""Rule #3 — DOCX extraction fidelity audit.

One-off script to inspect what actually got extracted from a DOCX file.
Prints each extraction source (paragraph, table, text box, header) with
provenance labels, and flags any XML artifacts in the extracted text.

Usage::

    python -m app.scripts.docx_extraction_audit --file path/to/document.docx
"""

from __future__ import annotations

import argparse
import io
import re
import sys
import logging

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)


def _has_xml_artifacts(text: str) -> bool:
    """Detect leftover XML/formatting artifacts that would push text
    semantically away from a clean natural-language query."""
    patterns = [
        r"<w:",           # raw WordprocessingML tags
        r"<a:",           # DrawingML
        r"<mc:",          # markup compatibility
        r"&#\d+;",        # numeric character references
        r"<[^>]+>",       # any HTML/XML tags
        r"\\x[0-9a-f]{2}",  # escaped hex bytes
    ]
    return any(re.search(p, text) for p in patterns)


def audit_docx(file_path: str) -> None:
    """Extract and audit a DOCX file, printing provenance for each text block."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed: pip install python-docx")
        sys.exit(2)

    with open(file_path, "rb") as f:
        content_bytes = f.read()

    doc = Document(io.BytesIO(content_bytes))

    print(f"\n{'='*70}")
    print(f"DOCX Extraction Audit: {file_path}")
    print(f"{'='*70}\n")

    total_chars = 0
    xml_artifact_count = 0

    # 1. Body paragraphs
    print("--- BODY PARAGRAPHS ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        has_artifacts = _has_xml_artifacts(text)
        flag = " [XML ARTIFACT!]" if has_artifacts else ""
        print(f"  [para:{i}]{flag} {text[:200]}{'...' if len(text) > 200 else ''}")
        total_chars += len(text)
        if has_artifacts:
            xml_artifact_count += 1

    # 2. Tables
    print("\n--- TABLES ---")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text.strip()
                    if not text:
                        continue
                    has_artifacts = _has_xml_artifacts(text)
                    flag = " [XML ARTIFACT!]" if has_artifacts else ""
                    print(
                        f"  [table:{t_idx}/row:{r_idx}/cell:{c_idx}/para:{p_idx}]{flag} "
                        f"{text[:200]}{'...' if len(text) > 200 else ''}"
                    )
                    total_chars += len(text)
                    if has_artifacts:
                        xml_artifact_count += 1

    # 3. Headers and footers
    print("\n--- HEADERS / FOOTERS ---")
    for s_idx, section in enumerate(doc.sections):
        for label, hf in [
            ("header", section.header),
            ("footer", section.footer),
        ]:
            if hf is None:
                continue
            try:
                for p_idx, para in enumerate(hf.paragraphs):
                    text = para.text.strip()
                    if not text:
                        continue
                    has_artifacts = _has_xml_artifacts(text)
                    flag = " [XML ARTIFACT!]" if has_artifacts else ""
                    print(
                        f"  [section:{s_idx}/{label}/para:{p_idx}]{flag} "
                        f"{text[:200]}{'...' if len(text) > 200 else ''}"
                    )
                    total_chars += len(text)
                    if has_artifacts:
                        xml_artifact_count += 1
            except Exception:
                pass

    # 4. Text boxes
    print("\n--- TEXT BOXES ---")
    try:
        body_xml = doc.element
        ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        tb_idx = 0
        for txbx in body_xml.iter(f"{ns_w}txbxContent"):
            for t_elem in txbx.iter(f"{ns_w}t"):
                text = (t_elem.text or "").strip()
                if not text:
                    continue
                has_artifacts = _has_xml_artifacts(text)
                flag = " [XML ARTIFACT!]" if has_artifacts else ""
                print(
                    f"  [textbox:{tb_idx}]{flag} "
                    f"{text[:200]}{'...' if len(text) > 200 else ''}"
                )
                total_chars += len(text)
                if has_artifacts:
                    xml_artifact_count += 1
                tb_idx += 1
    except Exception as e:
        print(f"  (text box extraction failed: {e})")

    # Summary
    print(f"\n{'='*70}")
    print(f"Total extracted chars: {total_chars}")
    if xml_artifact_count:
        print(f"WARNING: {xml_artifact_count} text blocks contain XML artifacts!")
        print("These artifacts push text semantically away from natural-language queries.")
    else:
        print("No XML artifacts detected — extraction looks clean.")
    print(f"{'='*70}\n")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Rule #3: DOCX extraction fidelity audit"
    )
    parser.add_argument("--file", required=True, help="Path to DOCX file")
    args = parser.parse_args()
    audit_docx(args.file)


if __name__ == "__main__":
    main()
